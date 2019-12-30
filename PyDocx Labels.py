from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
import pandas as pd
from math import ceil
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement
from tkinter.filedialog import askopenfilename
from tkinter import filedialog, messagebox, Tk, simpledialog
from os.path import basename
from os import getcwd

FrontTag = Document()
BackTag = Document()

testing = False

if testing:
    filename = 'VA - Nola.xlsx'
    ShipmentNumber = 'ShipmentNumber'
    ContainerNumber = 'ContainerNumber'
    save_dir = getcwd()

else:
    DialogueTitle = "Nadeau Print Label Generator"
    root = Tk()
    root.withdraw()
    messagebox.showinfo(DialogueTitle, 'Please select a Manifest to create labels.')
    filename = askopenfilename()

    ShipmentNumber = simpledialog.askstring(DialogueTitle, 'Internal Container Number')

    ContainerNumber = simpledialog.askstring(DialogueTitle, 'Vendor Container Number')

    messagebox.showinfo(DialogueTitle, 'Please select a save location for your files.')
    save_dir = filedialog.askdirectory()


Store = pd.read_excel(filename, skiprows=5)
Store = Store.dropna(how='any')
Store.columns = Store.columns.str.upper()

total_qty = 0
for total in range(0, len(Store['QTY']) - 1):
    total_qty += Store['QTY'][total]

pages = 50
pages = (total_qty/pages) + 1
cnt = 0

BackTagSections = BackTag.sections
FrontTagSections = FrontTag.sections

for BTS, FTS in zip(BackTagSections, FrontTagSections):
    BTS.top_margin = Inches(0.35)
    BTS.bottom_margin = Inches(0.35)
    BTS.left_margin = Inches(0.25)
    BTS.right_margin = Inches(0.25)

    FTS.top_margin = Inches(0.35)
    FTS.bottom_margin = Inches(0.35)
    FTS.right_margin = Inches(0.25)
    FTS.left_margin = Inches(0.25)


def prevent_document_break(document):
    tags = document.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]
        child = OxmlElement('w:cantSplit')
        tag.append(child)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def set_col_widths(table, width):
    widths = (Inches(width), Inches(width), Inches(width), Inches(width), Inches(width))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def front_label():
    qty = list(Store['QTY'])
    count = 0
    for page in range(0, ceil(pages)):
        table = FrontTag.add_table(rows=10, cols=5)
        table.allow_autofit = False
        table.BreakAcrossPages = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            set_col_widths(table, 1.65)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Inches(1)
            for cell in row.cells:
                try:
                    p = cell.add_paragraph(Store['ITEM'][count] + '\n')
                    p.add_run('$' + str(int(Store['PRICE'][count]))).bold = True
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except (IndexError, KeyError):
                    break

                if qty[count] > 1:
                    qty[count] -= 1
                else:
                    count += 1
    prevent_document_break(FrontTag)
    if '.xlsx' in filename:
        FrontTag.save(save_dir + '/' + basename(filename).replace('.xlsx', '') + ' Front Tag.docx')
    else:
        FrontTag.save(save_dir + '/' + basename(filename).replace('.xls', '') + ' Front Tag.docx')


def back_tag():
    qty = list(Store['QTY'])
    count = 0
    for page in range(0, ceil(pages)):
        table = BackTag.add_table(rows=10, cols=5)
        table.autofit = False
        table.BreakAcrossPages = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            # set_col_widths(table, 1.65)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Inches(1)
            for cell in row.cells:
                try:
                    p = cell.paragraphs[0].add_run(str(Store['ITEM'][count]))
                    p.bold = True
                    p.font.name = 'Arial Narrow'
                    p.font.size = Pt(12)

                    line = cell.paragraphs[0].add_run('\n' + ShipmentNumber)
                    line.font.size = Pt(12)
                    line.font.name = 'Arial Narrow'

                    line = cell.paragraphs[0].add_run('\n' + ContainerNumber)
                    line.font.size = Pt(12)
                    line.font.name = 'Arial Narrow'

                    line = cell.paragraphs[0].add_run('\n' + Store['DESCRIPTION'][count])
                    line.font.size = Pt(9)
                    line.font.name = 'Arial Narrow'
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                except (IndexError, KeyError):
                    break

                if qty[count] > 1:
                    qty[count] -= 1
                else:
                    count += 1
    prevent_document_break(BackTag)
    if '.xlsx' in filename:
        BackTag.save(save_dir + '/' + basename(filename).replace('.xlsx', '') + ' Back Tag.docx')
    else:
        BackTag.save(save_dir + '/' + basename(filename).replace('.xls', '') + ' Back Tag.docx')


front_label()
back_tag()

